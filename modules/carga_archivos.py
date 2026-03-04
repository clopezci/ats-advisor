# ==========================================================
#  ATS Advisor
#Herramienta tecnológica de análisis y mejora de postulaciones laborales
#
#Desarrollado por Carlos Emilio López (clopezci@hotmail.com)
#Proyecto independiente con propósito educativo y social
#Año: 2025-2026
# ----------------------------------------------------------
#  Descripción:
    #  ATS Advisor es una herramienta educativa. Evalúa
    #  la compatibilidad entre una hoja de vida (CV) y una oferta
    #  laboral, simulando el funcionamiento de un sistema ATS.
    #
    #  Propiedad Intelectual:
        #  © 2025-2026 Carlos Emilio López
        #  Licencia de uso: Código abierto con fines educativos,
        #  investigación, y mejora libre bajo reconocimiento de autoría.
        #
        #  Descargo de responsabilidad:
            #  Este software se proporciona "tal cual", sin garantía de
            #  precisión o adecuación comercial. El autor
            #  no se hacen responsables del uso indebido ni de decisiones
            #  tomadas con base en sus resultados. Los usuarios pueden
            #  modificar y adaptar el código respetando la autoría original.
            #
            #  Contacto:
                #  Carlos Emilio López - clopezci@hotmail.com
                # ==========================================================


# ==========================
# carga_archivos.py - Gestión de CVs y ofertas laborales (versión robusta)
# ==========================
# - Un solo root Tk oculto por llamada y destruido al final.
# - Manejo de excepciones en PDF/DOCX con mensajes de ayuda.
# - Lectura de tablas en DOCX (común en CVs).
# - Detección de PDFs vacíos (posible escaneado sin OCR).
# - Fallback por consola si el diálogo de Tk falla o el usuario cancela.
# - Soporte opcional .txt para pruebas.
# Autor: Carlos Emilio López (Proyecto TFM)
# ===========================================

import os
import sys
from contextlib import suppress

# Dependencias gráficas
with suppress(Exception):
    from tkinter import Tk, filedialog
    from tkinter.simpledialog import askstring

# Dependencias de documento
with suppress(Exception):
    import docx
with suppress(Exception):
    import PyPDF2


# ----------------------------
# Helpers de UI (Tk)
# ----------------------------
def _with_tk_dialogs(func):
    """Decora funciones que usan Tk: crea/destroye root de forma segura + lo trae al frente."""
    def wrapper(*args, **kwargs):
        root = None
        try:
            root = Tk()
            root.withdraw()

            #  Forzar al frente (Windows)
            root.lift()
            root.attributes("-topmost", True)
            root.update_idletasks()
            root.update()

        except Exception:
            root = None

        try:
                #  Pasar root a la función para usarlo como "parent"
                if root is not None:
                    root.after(200, lambda: root.attributes("-topmost", False))

                return func(*args, **kwargs, _tk_root=root)
        finally:
            if root is not None:
                try:
                    root.destroy()
                except Exception:
                    pass
    return wrapper


# ----------------------------
# CARGAR CV
# ----------------------------
@_with_tk_dialogs
def cargar_cv(_tk_root=None):
    """
    Abre una ventana para seleccionar CV (.pdf o .docx). Fallback por consola si falla.
    Retorna:
        str ruta del archivo o "" si no hay selección.
    """
    # Intento con Tk
    try:
        print("Seleccione su archivo de CV (.pdf o .docx) en la ventana emergente...")
        ruta = filedialog.askopenfilename(
            parent=_tk_root,  #  clave: sale encima
            title="Seleccionar archivo de CV",
            filetypes=[("PDF o Word", "*.pdf *.docx"), ("Todos", "*.*")]
        )
        if ruta:
            return ruta
    except Exception:
        pass

    # Fallback por consola
    print("No se pudo abrir el diálogo gráfico. Ingresa la ruta del CV (.pdf/.docx) o deja vacío para cancelar:")
    ruta = input("> ").strip('" ').strip()
    if not ruta:
        print("⚠️ No se seleccionó ningún archivo.")
        return ""
    if not os.path.exists(ruta):
        print("⚠️ La ruta no existe.")
        return ""
    return ruta


# ----------------------------
# CARGAR OFERTA
# ----------------------------
@_with_tk_dialogs
def cargar_oferta(_tk_root=None):
    """
    Ventana para pegar oferta (Toplevel propio, siempre al frente).
    Fallback por consola si falla.
    Retorna:
        str texto de la oferta ("" si se cancela).
    """
    try:
        import tkinter as tk
        from tkinter import ttk
        from tkinter.scrolledtext import ScrolledText

        print("Escriba o pegue el texto de la vacante en la ventana emergente.")

        #  Ventana propia (control total)
        win = tk.Toplevel(_tk_root) if _tk_root is not None else tk.Tk()
        win.title("Vacante - Pega el texto de la oferta")
        win.geometry("900x650")
        win.minsize(700, 450)

        #  Siempre al frente
        win.lift()
        win.attributes("-topmost", True)
        win.after(250, lambda: win.attributes("-topmost", False))
        win.focus_force()

        # Estilo
        try:
            ttk.Style(win).theme_use("clam")
        except Exception:
            pass

        ttk.Label(
            win,
            text="Pega aquí el texto completo de la oferta de empleo:",
            font=("Segoe UI", 11, "bold")
        ).pack(pady=(10, 6))

        box = ScrolledText(win, wrap="word", font=("Segoe UI", 10))
        box.pack(fill="both", expand=True, padx=12, pady=8)

        # Botonera
        btns = ttk.Frame(win)
        btns.pack(pady=(0, 12))

        result = {"texto": ""}

        def _usar():
            result["texto"] = (box.get("1.0", "end") or "").strip()
            win.destroy()

        def _cancelar():
            result["texto"] = ""
            win.destroy()

        ttk.Button(btns, text="Usar este texto", command=_usar).pack(side="left", padx=8)
        ttk.Button(btns, text="Cancelar", command=_cancelar).pack(side="left", padx=8)

        win.protocol("WM_DELETE_WINDOW", _cancelar)

        # Modal ligera para que no quede detrás
        try:
            win.grab_set()
        except Exception:
            pass

        win.wait_window()
        if result["texto"]:
            return result["texto"]

    except Exception:
        pass

    # Fallback por consola con EOF 
    print("No se pudo abrir el diálogo gráfico.")
    print("Pega el texto de la oferta y finaliza con una línea que contenga solo 'EOF':")
    lines = []
    while True:
        try:
            line = input()
        except EOFError:
            break
        if line.strip() == "EOF":
            break
        lines.append(line)
    return "\n".join(lines).strip()


# ----------------------------
# LEER CV COMO TEXTO
# ----------------------------
def leer_cv_como_texto(ruta):
    """
    Lee PDF/DOCX/TXT y devuelve texto plano. Maneja errores y casos sin texto.
    Retorna:
        str texto extraído ("" si no se pudo).
    """
    if not ruta:
        return ""

    ruta_lower = ruta.lower()
    try:
        if ruta_lower.endswith(".pdf"):
            return extraer_texto_pdf(ruta)
        elif ruta_lower.endswith(".docx"):
            return extraer_texto_docx(ruta)
        elif ruta_lower.endswith(".txt"):
            return extraer_texto_txt(ruta)
        else:
            raise ValueError("Formato no compatible. Usa .pdf, .docx o .txt")
    except Exception as e:
        print(f"❌ Error al leer CV: {e}")
        return ""


# ----------------------------
# EXTRAER TEXTO DE TXT 
# ----------------------------
def extraer_texto_txt(ruta):
    try:
        with open(ruta, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise RuntimeError(f"No se pudo leer TXT: {e}") from e


# ----------------------------
# EXTRAER TEXTO DE DOCX (incluye tablas)
# ----------------------------
def extraer_texto_docx(ruta):
    if 'docx' not in sys.modules:
        raise RuntimeError("La librería 'python-docx' no está instalada. Ejecuta: pip install python-docx")

    try:
        d = docx.Document(ruta)
    except Exception as e:
        raise RuntimeError(f"No se pudo abrir el DOCX: {e}") from e

    partes = []

    # Párrafos
    for p in d.paragraphs:
        if p.text and p.text.strip():
            partes.append(p.text.strip())

    # Tablas
    for tbl in d.tables:
        for row in tbl.rows:
            celdas = []
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if txt:
                    celdas.append(txt)
            if celdas:
                partes.append(" | ".join(celdas))

    texto = "\n".join(partes).strip()
    if not texto:
        raise RuntimeError("El DOCX parece no contener texto legible.")
    return texto


# ----------------------------
# EXTRAER TEXTO DE PDF (detección escaneados)
# ----------------------------
def extraer_texto_pdf(ruta):
    if 'PyPDF2' not in sys.modules:
        raise RuntimeError("La librería 'PyPDF2' no está instalada. Ejecuta: pip install PyPDF2")

    try:
        texto = []
        with open(ruta, "rb") as archivo:
            lector = PyPDF2.PdfReader(archivo)
            if lector.is_encrypted:
                # intento de desencriptado vacío (muchos PDFs permiten con cadena vacía)
                with suppress(Exception):
                    lector.decrypt("")
            for i, pagina in enumerate(lector.pages):
                try:
                    page_text = pagina.extract_text() or ""
                except Exception:
                    page_text = ""
                texto.append(page_text)
        joined = "\n".join(texto).strip()
        if not joined:
            
            raise RuntimeError(
                "El PDF no contiene texto extraíble (posible escaneado o protegido). "
                "Convierte a PDF con texto (OCR) o exporta a DOCX/TXT antes de analizar."
            )
        return joined
    except Exception as e:
        raise RuntimeError(f"No se pudo leer el PDF: {e}") from e
